# -----------------------------------------------------------------------------
# DOCX EXPORT: Pretvara HTML iz generatora u pravi .docx Word dokument
# Koristi python-docx za nativni Word format (radi na mobitelu, tabletu, PC-u)
# -----------------------------------------------------------------------------
import io
import re
from html.parser import HTMLParser
from html import unescape

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# --- Konstante formatiranja za hrvatske pravne dokumente ---
FONT_NAME = "Times New Roman"
FONT_SIZE_BODY = Pt(12)
FONT_SIZE_HEADER = Pt(14)
FONT_SIZE_SECTION = Pt(11)
FONT_SIZE_SMALL = Pt(10)
LINE_SPACING = 1.15
MARGIN_CM = 2.5


class _HtmlToDocxParser(HTMLParser):
    """
    Parsira HTML iz nasih generatora i pretvara u python-docx elemente.
    Podrzava: div klase (header-doc, party-info, doc-body, justified,
    section-title, cost-table, signature-row, clausula), b, i, u, br,
    table/tr/td, ul/ol/li, span.
    """

    def __init__(self, doc):
        super().__init__()
        self.doc = doc
        self.current_paragraph = None
        self.current_run_props = {
            'bold': False,
            'italic': False,
            'underline': False,
            'size': FONT_SIZE_BODY,
        }
        self.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        self.tag_stack = []
        self.class_stack = []
        self.in_table = False
        self.table_rows = []
        self.current_row = []
        self.current_cell_text = ""
        self.current_cell_bold = False
        self.current_cell_align = "left"
        self.in_list = False
        self.list_type = "ul"
        self.list_counter = 0
        self.skip_content = False
        self.pending_text = ""

    def _flush_text(self):
        """Zapisuje sakupljeni tekst u trenutni paragraf."""
        if self.pending_text and self.current_paragraph and not self.in_table:
            run = self.current_paragraph.add_run(self.pending_text)
            run.font.name = FONT_NAME
            run.font.size = self.current_run_props['size']
            run.bold = self.current_run_props['bold']
            run.italic = self.current_run_props['italic']
            run.underline = self.current_run_props['underline']
            self.pending_text = ""

    def _new_paragraph(self, alignment=None):
        """Kreira novi paragraf s zadanim poravnanjem."""
        self._flush_text()
        p = self.doc.add_paragraph()
        p.alignment = alignment if alignment is not None else self.alignment
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.line_spacing = LINE_SPACING
        self.current_paragraph = p
        return p

    def _add_horizontal_line(self):
        """Dodaje horizontalnu liniju (bottom border na praznom paragrafu)."""
        self._flush_text()
        p = self.doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.space_before = Pt(6)
        pPr = p._element.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '000000')
        pBdr.append(bottom)
        pPr.append(pBdr)
        self.current_paragraph = None

    def _extract_classes(self, attrs):
        """Izvlaci CSS klase iz atributa."""
        for name, value in attrs:
            if name == "class":
                return value.split()
        return []

    def _extract_style(self, attrs):
        """Izvlaci inline style iz atributa."""
        for name, value in attrs:
            if name == "style":
                return value
        return ""

    def _extract_attr(self, attrs, attr_name):
        for name, value in attrs:
            if name == attr_name:
                return value
        return ""

    def handle_starttag(self, tag, attrs):
        classes = self._extract_classes(attrs)
        style = self._extract_style(attrs)
        self.tag_stack.append(tag)
        self.class_stack.append(classes)

        if tag == "div":
            if "header-doc" in classes:
                self._flush_text()
                p = self._new_paragraph(WD_ALIGN_PARAGRAPH.CENTER)
                self.current_run_props['bold'] = True
                self.current_run_props['size'] = FONT_SIZE_HEADER
            elif "section-title" in classes:
                self._flush_text()
                p = self._new_paragraph(WD_ALIGN_PARAGRAPH.LEFT)
                style_str = self._extract_style(attrs)
                if "text-align: center" in style_str:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                self.current_run_props['bold'] = True
                self.current_run_props['size'] = FONT_SIZE_SECTION
            elif "party-info" in classes:
                self._flush_text()
                self._new_paragraph(WD_ALIGN_PARAGRAPH.LEFT)
                self.current_run_props['bold'] = False
                self.current_run_props['size'] = FONT_SIZE_BODY
            elif "doc-body" in classes or "justified" in classes:
                self._flush_text()
                p = self._new_paragraph(WD_ALIGN_PARAGRAPH.JUSTIFY)
                self.current_run_props['bold'] = False
                self.current_run_props['size'] = FONT_SIZE_BODY
                if "clausula" in classes:
                    self.current_run_props['italic'] = True
            elif "signature-row" in classes:
                self._flush_text()
                # Prazan prostor prije potpisa
                for _ in range(2):
                    self.doc.add_paragraph()
            elif "signature-block" in classes:
                self._flush_text()
                self._add_horizontal_line()
                self._new_paragraph(WD_ALIGN_PARAGRAPH.CENTER)
            elif "cost-table" in classes:
                pass  # Handled by table tag
            else:
                # Genericki div - provjeri style za text-align
                if "text-align: center" in style or "text-align:center" in style:
                    self._flush_text()
                    self._new_paragraph(WD_ALIGN_PARAGRAPH.CENTER)
                    if "font-weight: bold" in style or "font-weight:bold" in style:
                        self.current_run_props['bold'] = True
                    if "font-size: 14px" in style or "font-size:14px" in style:
                        self.current_run_props['size'] = FONT_SIZE_HEADER
                elif "text-align: right" in style or "text-align:right" in style:
                    self._flush_text()
                    self._new_paragraph(WD_ALIGN_PARAGRAPH.RIGHT)
                elif "font-weight: bold" in style or "font-weight:bold" in style:
                    self._flush_text()
                    self._new_paragraph(WD_ALIGN_PARAGRAPH.LEFT)
                    self.current_run_props['bold'] = True
                    if "font-size: 14px" in style:
                        self.current_run_props['size'] = FONT_SIZE_HEADER
                elif "border: 2px solid" in style:
                    # Okvir za rjesenje - dodaj prazan red i naglaseni paragraf
                    self._flush_text()
                    self.doc.add_paragraph()
                elif style and self.current_paragraph is None:
                    self._new_paragraph(WD_ALIGN_PARAGRAPH.LEFT)

        elif tag == "b":
            self._flush_text()
            self.current_run_props['bold'] = True

        elif tag == "i":
            self._flush_text()
            self.current_run_props['italic'] = True

        elif tag == "u":
            self._flush_text()
            self.current_run_props['underline'] = True

        elif tag == "br":
            if self.in_table:
                self.current_cell_text += "\n"
            elif self.current_paragraph:
                self._flush_text()
                # Novi paragraf (^p) umjesto manual line break-a (^l / w:br)
                # w:br stvara ^l koji zahtijeva rucni replace u Wordu
                old_p = self.current_paragraph
                p = self.doc.add_paragraph()
                p.alignment = old_p.alignment
                p.paragraph_format.space_after = old_p.paragraph_format.space_after
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.line_spacing = old_p.paragraph_format.line_spacing
                self.current_paragraph = p

        elif tag == "hr":
            self._add_horizontal_line()

        elif tag == "span":
            style_str = self._extract_style(attrs)
            if "font-weight: normal" in style_str or "font-weight:normal" in style_str:
                self._flush_text()
                self.current_run_props['bold'] = False
            if "font-size: 10pt" in style_str or "font-size:10pt" in style_str:
                self._flush_text()
                self.current_run_props['size'] = FONT_SIZE_SMALL
            if "font-size: 11pt" in style_str or "font-size:11pt" in style_str:
                self._flush_text()
                self.current_run_props['size'] = FONT_SIZE_SECTION
            if "font-size: 12pt" in style_str:
                self._flush_text()
                self.current_run_props['size'] = FONT_SIZE_BODY

        elif tag == "table":
            self._flush_text()
            self.in_table = True
            self.table_rows = []

        elif tag == "tr":
            if self.in_table:
                self.current_row = []
                self.current_cell_bold = False

        elif tag == "td":
            if self.in_table:
                self.current_cell_text = ""
                align = self._extract_attr(attrs, "align")
                self.current_cell_align = align or "left"
                style_str = self._extract_style(attrs)
                self.current_cell_bold = "font-weight: bold" in style_str or "font-weight:bold" in style_str

        elif tag == "ul":
            self.in_list = True
            self.list_type = "ul"
            self.list_counter = 0

        elif tag == "ol":
            self.in_list = True
            self.list_type = "ol"
            self.list_counter = 0

        elif tag == "li":
            self._flush_text()
            self.list_counter += 1
            p = self._new_paragraph(WD_ALIGN_PARAGRAPH.LEFT)
            p.paragraph_format.left_indent = Cm(1)
            if self.list_type == "ol":
                prefix = f"{self.list_counter}. "
            else:
                prefix = "\u2022 "
            run = p.add_run(prefix)
            run.font.name = FONT_NAME
            run.font.size = FONT_SIZE_BODY

    def handle_endtag(self, tag):
        if self.tag_stack and self.tag_stack[-1] == tag:
            self.tag_stack.pop()
            classes = self.class_stack.pop() if self.class_stack else []
        else:
            classes = []

        if tag == "div":
            self._flush_text()
            if "header-doc" in classes:
                self.current_run_props['bold'] = False
                self.current_run_props['size'] = FONT_SIZE_BODY
            elif "section-title" in classes:
                self.current_run_props['bold'] = False
                self.current_run_props['size'] = FONT_SIZE_BODY
            elif "clausula" in classes or ("doc-body" in classes and "clausula" in classes):
                self.current_run_props['italic'] = False
            # Reset za sve divove
            self.current_run_props['bold'] = False
            self.current_run_props['size'] = FONT_SIZE_BODY

        elif tag == "b":
            self._flush_text()
            self.current_run_props['bold'] = False

        elif tag == "i":
            self._flush_text()
            self.current_run_props['italic'] = False

        elif tag == "u":
            self._flush_text()
            self.current_run_props['underline'] = False

        elif tag == "table":
            if self.in_table and self.table_rows:
                self._create_table()
            self.in_table = False
            self.table_rows = []

        elif tag == "tr":
            if self.in_table and self.current_row:
                self.table_rows.append(self.current_row)
                self.current_row = []

        elif tag == "td":
            if self.in_table:
                self.current_row.append({
                    'text': self.current_cell_text.strip(),
                    'bold': self.current_cell_bold,
                    'align': self.current_cell_align,
                })
                self.current_cell_text = ""

        elif tag in ("ul", "ol"):
            self.in_list = False
            self.list_counter = 0

    def handle_data(self, data):
        text = data
        if not text:
            return

        if self.in_table:
            self.current_cell_text += text
            return

        if self.current_paragraph is None:
            # Stvori paragraf ako ga nema i tekst nije prazan
            cleaned = text.strip()
            if cleaned:
                self._new_paragraph(WD_ALIGN_PARAGRAPH.LEFT)
                self.pending_text += cleaned
        else:
            self.pending_text += text

    def handle_entityref(self, name):
        char = unescape(f"&{name};")
        self.handle_data(char)

    def handle_charref(self, name):
        char = unescape(f"&#{name};")
        self.handle_data(char)

    def _create_table(self):
        """Kreira Word tablicu iz sakupljenih redova."""
        if not self.table_rows:
            return

        max_cols = max(len(row) for row in self.table_rows)
        if max_cols == 0:
            return

        table = self.doc.add_table(rows=len(self.table_rows), cols=max_cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        for i, row_data in enumerate(self.table_rows):
            for j, cell_data in enumerate(row_data):
                if j < max_cols:
                    cell = table.rows[i].cells[j]
                    p = cell.paragraphs[0]
                    run = p.add_run(cell_data['text'])
                    run.font.name = FONT_NAME
                    run.font.size = FONT_SIZE_BODY
                    run.bold = cell_data['bold']

                    if cell_data['align'] == "right":
                        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    elif cell_data['align'] == "center":
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    else:
                        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    def finalize(self):
        """Zavrsava parsiranje - flusha preostali tekst."""
        self._flush_text()


def _dodaj_broj_stranice(section):
    """Dodaje 'Stranica X od Y' u footer dokumenta (desno poravnato)."""
    footer = section.footer
    footer.is_linked_to_previous = False
    paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    run = paragraph.add_run("Stranica ")
    run.font.name = FONT_NAME
    run.font.size = FONT_SIZE_SMALL
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    # PAGE field
    fldChar1 = run._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'begin'})
    run2 = paragraph.add_run()
    run2._element.append(fldChar1)
    instrText = run2._element.makeelement(qn('w:instrText'), {})
    instrText.text = " PAGE "
    run3 = paragraph.add_run()
    run3._element.append(instrText)
    fldChar2 = run3._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'end'})
    run4 = paragraph.add_run()
    run4._element.append(fldChar2)

    run5 = paragraph.add_run(" od ")
    run5.font.name = FONT_NAME
    run5.font.size = FONT_SIZE_SMALL
    run5.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    # NUMPAGES field
    fldChar3 = run5._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'begin'})
    run6 = paragraph.add_run()
    run6._element.append(fldChar3)
    instrText2 = run6._element.makeelement(qn('w:instrText'), {})
    instrText2.text = " NUMPAGES "
    run7 = paragraph.add_run()
    run7._element.append(instrText2)
    fldChar4 = run7._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'end'})
    run8 = paragraph.add_run()
    run8._element.append(fldChar4)


def _dodaj_watermark(section, tekst="NACRT"):
    """Dodaje dijagonalni watermark tekst u header dokumenta (raw XML VML shape)."""
    from lxml import etree

    header = section.header
    header.is_linked_to_previous = False
    paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()

    # VML watermark kao raw XML - jedini pouzdani nacin s python-docx
    watermark_xml = (
        '<w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        '<w:pict>'
        '<v:shapetype xmlns:v="urn:schemas-microsoft-com:vml" id="_x0000_t136" '
        'coordsize="21600,21600" path="m@7,l@8,m@5,21600l@6,21600e"/>'
        '<v:shape xmlns:v="urn:schemas-microsoft-com:vml" '
        'xmlns:o="urn:schemas-microsoft-com:office:office" '
        'id="PowerPlusWaterMarkObject" type="#_x0000_t136" '
        'style="position:absolute;margin-left:0;margin-top:0;width:450pt;height:100pt;'
        'rotation:315;z-index:-251658752;'
        'mso-position-horizontal:center;mso-position-horizontal-relative:margin;'
        'mso-position-vertical:center;mso-position-vertical-relative:margin" '
        f'fillcolor="#d0d0d0" stroked="f">'
        f'<v:textpath style="font-family:&quot;Times New Roman&quot;;font-size:1pt" '
        f'string="{tekst}"/>'
        '<v:fill opacity=".25"/>'
        '</v:shape>'
        '</w:pict>'
        '</w:r>'
    )

    run_element = etree.fromstring(watermark_xml)
    paragraph._element.append(run_element)


def _dodaj_header_naslov(section, naslov):
    """Dodaje naziv dokumenta u header (lijevo poravnato, sivi tekst)."""
    header = section.header
    header.is_linked_to_previous = False
    # Ako vec postoji paragraf (od watermark-a), dodaj novi
    if header.paragraphs and header.paragraphs[0].text.strip():
        paragraph = header.add_paragraph()
    else:
        paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()

    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(naslov)
    run.font.name = FONT_NAME
    run.font.size = FONT_SIZE_SMALL
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    run.font.italic = True

    # Dodaj tanku liniju ispod headera
    pPr = paragraph._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), 'CCCCCC')
    pBdr.append(bottom)
    pPr.append(pBdr)


def html_u_docx(
    html_sadrzaj,
    watermark=None,
    naslov_dokumenta=None,
    user_id=None,
    doc_type=None,
    plan=None,
    input_dict=None,
    generator_module_path=None,
):
    """
    Pretvara HTML iz generatora u pravi .docx Word dokument.
    Vraca bytes spremne za download.

    Args:
        html_sadrzaj: HTML string iz generator/* funkcije
        watermark: tekst watermark-a (npr. "NACRT", "DRAFT") ili None — sredisnji vodeni zig
        naslov_dokumenta: naziv dokumenta za header ili None
        user_id: K3 forenzicki audit — Supabase user_id (UUID); None za guest
        doc_type: K3 forenzicki audit — npr. 'tuzba', 'ovrha' (sluzi u download_log)
        plan: 'free' ili 'pro'; ako None, citamo iz entitlements (graceful degrade ako Supabase nije konfiguriran)

    K3 watermark integracija (per-doc forensic serial):
        - Generira jedinstveni 'NN-NNNN-NNNNNN' broj
        - Visible footer: "Generirano iz LegalTechSuite Pro - ID: NN-NNNN-NNNNNN" (free)
                          ili "ID: NN-NNNN-NNNNNN" (pro, cleaner)
        - Invisible XML metadata: dc:identifier (forenzicki alati to vide)
        - Tihi fail: ako watermark/entitlements moduli nisu dostupni (npr. testovi
          koji ne importiraju watermark.py), docx se generira bez K3 sloja.
    """
    doc = Document()

    # Postavi margine (2.5 cm - standard za pravne dokumente)
    for section in doc.sections:
        section.top_margin = Cm(MARGIN_CM)
        section.bottom_margin = Cm(MARGIN_CM)
        section.left_margin = Cm(MARGIN_CM)
        section.right_margin = Cm(MARGIN_CM)
        # Dodaj broj stranice u footer
        _dodaj_broj_stranice(section)
        # Watermark (sredisnji "NACRT" tekst, neovisan o K3 forensic serial)
        if watermark:
            _dodaj_watermark(section, watermark)
        # Header s nazivom dokumenta
        if naslov_dokumenta:
            _dodaj_header_naslov(section, naslov_dokumenta)

    # Postavi default font
    style = doc.styles['Normal']
    font = style.font
    font.name = FONT_NAME
    font.size = FONT_SIZE_BODY
    pf = style.paragraph_format
    pf.line_spacing = LINE_SPACING
    pf.space_after = Pt(2)

    # ─── K3: per-doc forensic serial broj (apply_watermark) ───────────────────
    # Generira serial PRIJE parsing-a tako da footer HTML moze biti dodan u
    # rendered HTML; XML metadata se utiska u doc objekt direktno.
    serial = None
    serial_h = None
    effective_plan = plan or "free"
    try:
        import watermark as _wm
        # Ako plan nije eksplicitno proslijeden, ucitaj iz entitlements (graceful)
        if plan is None:
            try:
                import entitlements as _ent
                effective_plan = "pro" if _ent.is_pro(user_id=user_id) else "free"
            except Exception:
                effective_plan = "free"
        # Generiraj serial + utiskiraj XML metadata (dc:identifier)
        if doc_type:
            serial, serial_h = _wm.apply_watermark(
                doc=doc,
                user_id=user_id or "guest",
                doc_type=doc_type,
                plan=effective_plan,
            )
            # Dodaj visible footer u HTML PRIJE parsiranja
            html_sadrzaj = html_sadrzaj + _wm.footer_html(serial, plan=effective_plan)
    except ImportError:
        # watermark modul nije dostupan (npr. minimalni test env) - tihi fall-through
        pass

    # Parsiraj HTML
    parser = _HtmlToDocxParser(doc)
    # Ocisti HTML prije parsiranja
    clean_html = html_sadrzaj.replace('&nbsp;', ' ')
    parser.feed(clean_html)
    parser.finalize()

    # Spremi u BytesIO
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    output_bytes = buffer.getvalue()

    # ─── K1: audit chain + K3 record download (tihi fail) ───────────────────
    if serial_h and doc_type:
        try:
            import entitlements as _ent
            audit_payload = None
            # K1 audit chain: aktivira se SAMO ako pozivatelj proslijedi
            # input_dict + generator_module_path (backwards compat: postojeci
            # generatori koji ne proslijede ostaju bez audit chain-a, samo K3 watermark)
            if input_dict is not None and generator_module_path is not None:
                try:
                    import audit_chain as _ac
                    parent = _ent.get_last_chain_hash(user_id=user_id)
                    audit_payload = _ac.compute_full_audit(
                        input_dict=input_dict,
                        output_bytes=output_bytes,
                        generator_module_path=generator_module_path,
                        parent_hash=parent,
                    )
                except Exception:
                    audit_payload = None  # tihi fall-through, K3 watermark ostaje
            _ent.record_download(
                doc_type=doc_type,
                doc_subtype=naslov_dokumenta or doc_type,
                serial_hash=serial_h,
                plan=effective_plan,
                user_id=user_id,
                audit=audit_payload,
            )
        except Exception:
            pass

    return output_bytes


def pripremi_za_docx(
    html_sadrzaj,
    watermark=None,
    naslov_dokumenta=None,
    user_id=None,
    doc_type=None,
    plan=None,
    input_dict=None,
    generator_module_path=None,
):
    """
    Glavna funkcija za export - zamjena za stari pripremi_za_word().
    Vraca bytes.

    K3 forensic watermark se aktivira automatski ako je `doc_type` proslijeden.
    K1 audit chain se aktivira ako su `input_dict` i `generator_module_path`
    proslijedjeni — racuna se input_canonical_hash, output_sha256,
    generator_version_hash i hash chain link, upisuje u download_log.
    user_id, plan se citaju iz entitlements ako nisu eksplicitno postavljeni.
    Backwards-compat: stari call `pripremi_za_docx(html, "NACRT", "Tuzba")` radi
    bez izmjene — K3 i K1 su no-op kad opcijski parametri nisu proslijedjeni.
    """
    return html_u_docx(
        html_sadrzaj,
        watermark=watermark,
        naslov_dokumenta=naslov_dokumenta,
        user_id=user_id,
        doc_type=doc_type,
        plan=plan,
        input_dict=input_dict,
        generator_module_path=generator_module_path,
    )
