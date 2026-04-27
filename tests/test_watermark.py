"""Tests za K3 watermark — per-doc serial generator + footer + XML metadata."""
import re

import pytest


def test_serial_format_matches_pattern():
    """Serial je 'NN-NNNN-NNNNNN' uppercase hex format."""
    import watermark as wm
    s = wm.generate_serial("user-1", "tuzba", timestamp_iso="2026-04-27T10:00:00+00:00", nonce="x")
    assert re.match(r"^[A-F0-9]{2}-[A-F0-9]{4}-[A-F0-9]{6}$", s), f"format wrong: {s}"


def test_serial_deterministic_for_fixed_inputs():
    """Isti input + nonce -> isti output (idempotent test-able)."""
    import watermark as wm
    a = wm.generate_serial("u1", "tuzba", timestamp_iso="2026-04-27T10:00:00+00:00", nonce="abc")
    b = wm.generate_serial("u1", "tuzba", timestamp_iso="2026-04-27T10:00:00+00:00", nonce="abc")
    assert a == b


def test_serial_unique_for_different_users():
    import watermark as wm
    a = wm.generate_serial("u1", "tuzba", timestamp_iso="2026-04-27T10:00:00+00:00", nonce="x")
    b = wm.generate_serial("u2", "tuzba", timestamp_iso="2026-04-27T10:00:00+00:00", nonce="x")
    assert a != b


def test_serial_unique_for_different_doc_types():
    import watermark as wm
    a = wm.generate_serial("u1", "tuzba", timestamp_iso="2026-04-27T10:00:00+00:00", nonce="x")
    b = wm.generate_serial("u1", "ovrha", timestamp_iso="2026-04-27T10:00:00+00:00", nonce="x")
    assert a != b


def test_random_nonce_makes_production_calls_unique():
    """Bez explicit nonce, secrets.token_hex(8) garantira unique."""
    import watermark as wm
    a = wm.generate_serial("u1", "tuzba")
    b = wm.generate_serial("u1", "tuzba")
    assert a != b


def test_serial_hash_is_full_sha256():
    import watermark as wm
    s = "AB-CDEF-123456"
    h = wm.serial_hash(s)
    assert len(h) == 64
    assert re.match(r"^[a-f0-9]{64}$", h)
    # Idempotent
    assert wm.serial_hash(s) == h


def test_footer_html_free_includes_brand_and_serial():
    import watermark as wm
    html = wm.footer_html("AB-CDEF-123456", plan="free")
    assert "LegalTechSuite Pro" in html
    assert "AB-CDEF-123456" in html


def test_footer_html_pro_omits_brand():
    """PRO tier dobije cleaner footer (samo ID, bez brand reklame)."""
    import watermark as wm
    html = wm.footer_html("AB-CDEF-123456", plan="pro")
    assert "LegalTechSuite Pro" not in html
    assert "AB-CDEF-123456" in html


def test_stamp_core_properties_sets_identifier():
    """Invisible XML metadata watermark u dc:identifier."""
    pytest.importorskip("docx")
    from docx import Document
    import watermark as wm
    doc = Document()
    wm.stamp_core_properties(doc, serial="AB-CDEF-123456")
    assert doc.core_properties.identifier == "AB-CDEF-123456"


def test_apply_watermark_returns_serial_and_hash():
    pytest.importorskip("docx")
    from docx import Document
    import watermark as wm
    doc = Document()
    serial, h = wm.apply_watermark(doc, user_id="u1", doc_type="tuzba", plan="pro")
    assert re.match(r"^[A-F0-9]{2}-[A-F0-9]{4}-[A-F0-9]{6}$", serial)
    assert len(h) == 64
    assert doc.core_properties.identifier == serial


def test_apply_watermark_handles_missing_user_id():
    """Guest user (user_id=''): default 'guest' fallback, ne padaj."""
    pytest.importorskip("docx")
    from docx import Document
    import watermark as wm
    doc = Document()
    serial, h = wm.apply_watermark(doc, user_id="", doc_type="tuzba", plan="free")
    assert serial  # ne prazan
    assert h


def test_apply_watermark_silent_on_non_docx():
    """Ako prosijedimo ne-Document objekt, stamp_core_properties ne pada."""
    import watermark as wm
    fake_doc = object()
    serial, h = wm.apply_watermark(fake_doc, user_id="u1", doc_type="tuzba")
    # Serial je vracen; XML stamping je tihi no-op
    assert serial and h
