# -----------------------------------------------------------------------------
# UNIT TESTOVI: docx_export.py
# Pokretanje: python -m pytest tests/test_docx_export.py -v
# -----------------------------------------------------------------------------
import sys
import os
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from docx_export import pripremi_za_docx, html_u_docx
from io import BytesIO
from docx import Document


class TestPripremiZaDocx:
    def test_returns_bytes(self):
        result = pripremi_za_docx("<div class='doc-body'>Test</div>")
        assert isinstance(result, bytes)

    def test_valid_docx(self):
        result = pripremi_za_docx("<div class='doc-body'>Test dokument</div>")
        doc = Document(BytesIO(result))
        assert len(doc.paragraphs) > 0

    def test_contains_text(self):
        result = pripremi_za_docx("<div class='doc-body'>Testni sadrzaj 12345</div>")
        doc = Document(BytesIO(result))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Testni sadrzaj 12345" in full_text

    def test_bold_text(self):
        result = pripremi_za_docx("<b>Bold tekst</b>")
        doc = Document(BytesIO(result))
        found_bold = False
        for p in doc.paragraphs:
            for run in p.runs:
                if run.bold and "Bold" in run.text:
                    found_bold = True
        assert found_bold

    def test_italic_text(self):
        result = pripremi_za_docx("<i>Italic tekst</i>")
        doc = Document(BytesIO(result))
        found_italic = False
        for p in doc.paragraphs:
            for run in p.runs:
                if run.italic and "Italic" in run.text:
                    found_italic = True
        assert found_italic

    def test_table(self):
        html = "<table><tr><td>Cell 1</td><td>Cell 2</td></tr></table>"
        result = pripremi_za_docx(html)
        doc = Document(BytesIO(result))
        assert len(doc.tables) >= 1

    def test_list(self):
        html = "<ul><li>Item 1</li><li>Item 2</li></ul>"
        result = pripremi_za_docx(html)
        doc = Document(BytesIO(result))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Item 1" in full_text

    def test_header_doc_class(self):
        html = "<div class='header-doc'>NASLOV DOKUMENTA</div>"
        result = pripremi_za_docx(html)
        doc = Document(BytesIO(result))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "NASLOV DOKUMENTA" in full_text

    def test_empty_input(self):
        result = pripremi_za_docx("")
        assert isinstance(result, bytes)

    def test_complex_document(self):
        """Test sa slozenim HTML-om koji simulira pravi generator output."""
        html = """
        <div class='header-doc'>TUŽBA</div>
        <div class='party-info'>
            <b>TUŽITELJ:</b> Test Ime<br>OIB: 12345678901
        </div>
        <div class='doc-body'>
            <div class='section-title'>I. ČINJENICE</div>
            <p>Tekst cinjenica s <b>boldanim</b> i <i>italic</i> dijelovima.</p>
            <table class='cost-table'>
                <tr><td>Stavka</td><td>100,00 EUR</td></tr>
            </table>
        </div>
        <div class='signature-row'>
            <div class='signature-block'>Potpis</div>
        </div>
        """
        result = pripremi_za_docx(html)
        doc = Document(BytesIO(result))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "TUŽBA" in full_text
        assert len(doc.tables) >= 1

    def test_page_numbers_in_footer(self):
        """Provjera da DOCX sadrzi footer (sekcija s brojem stranica)."""
        result = pripremi_za_docx("<div class='doc-body'>Test</div>")
        doc = Document(BytesIO(result))
        section = doc.sections[0]
        footer = section.footer
        # Footer should exist and have paragraphs
        assert footer is not None
        assert len(footer.paragraphs) > 0
